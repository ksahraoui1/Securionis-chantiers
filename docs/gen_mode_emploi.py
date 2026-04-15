#!/usr/bin/env python3
"""Génère le mode d'emploi Securionis Chantiers en .docx"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Couleurs ──────────────────────────────────────────────────────────────────
BLUE       = RGBColor(0x1E, 0x3A, 0x5F)
BLUE_LIGHT = RGBColor(0x2E, 0x75, 0xB6)
BLUE_BG    = "D6E4F0"
GREEN      = RGBColor(0x1F, 0x6B, 0x3A)
GREEN_BG   = "D6EFD8"
RED        = RGBColor(0xB9, 0x1C, 0x1C)
RED_BG     = "FEE2E2"
ORANGE     = RGBColor(0xB4, 0x53, 0x09)
ORANGE_BG  = "FEF3C7"
GREY_BG    = "F3F4F6"
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
DARK       = RGBColor(0x33, 0x33, 0x33)

# ─── Helpers XML ───────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)

def set_cell_borders(cell, color="CBD5E1"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcBorders)

def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    existing = tcPr.find(qn('w:tcMar'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcMar)

def set_col_width(cell, width_dxa):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_dxa))
    tcW.set(qn('w:type'), 'dxa')
    existing = tcPr.find(qn('w:tcW'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcW)

def set_table_width(table, width_dxa):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(width_dxa))
    tblW.set(qn('w:type'), 'dxa')
    existing = tblPr.find(qn('w:tblW'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblW)

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)

def add_paragraph_spacing(para, before=0, after=60):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    existing = pPr.find(qn('w:spacing'))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(spacing)

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_breaks.WD_BREAK.PAGE)

# ─── Constructeurs ─────────────────────────────────────────────────────────────
def styled_cell(table_cell, text, bg=None, bold=False, font_color=None,
                font_size=10, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, col_width=None):
    table_cell.text = ""
    set_cell_borders(table_cell)
    set_cell_margins(table_cell)
    if bg:
        set_cell_bg(table_cell, bg)
    if col_width:
        set_col_width(table_cell, col_width)
    p = table_cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    if font_color:
        run.font.color.rgb = font_color
    return table_cell

def header_cell(tc, text, col_width=None):
    return styled_cell(tc, text, bg="1E3A5F", bold=True, font_color=WHITE,
                       font_size=10, col_width=col_width)

def grey_label_cell(tc, text, col_width=None):
    return styled_cell(tc, text, bg=GREY_BG, bold=True, font_size=9.5, col_width=col_width)

def value_cell(tc, text, col_width=None):
    return styled_cell(tc, text, font_size=9.5, col_width=col_width)

def badge_cell(tc, text, bg, color_rgb, col_width=None):
    return styled_cell(tc, text, bg=bg, bold=True, font_color=color_rgb,
                       font_size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER, col_width=col_width)

# ─── Document ──────────────────────────────────────────────────────────────────
doc = Document()

# Marges de page (2cm)
from docx.oxml import OxmlElement
section = doc.sections[0]
section.page_width  = Cm(21.0)
section.page_height = Cm(29.7)
section.left_margin = section.right_margin = Cm(2.5)
section.top_margin  = section.bottom_margin = Cm(2.0)

# Styles de base
styles = doc.styles
normal = styles['Normal']
normal.font.name = "Arial"
normal.font.size = Pt(10.5)

# ─── Helpers de paragraphes ────────────────────────────────────────────────────
def add_h1(doc, text):
    p = doc.add_paragraph()
    add_paragraph_spacing(p, before=240, after=120)
    pPr = p._p.get_or_add_pPr()
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), 'Heading1')
    pPr.insert(0, pStyle)
    # Fond bleu via shading du paragraphe non disponible directement — on stylise le run
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = BLUE
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    add_paragraph_spacing(p, before=160, after=80)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = BLUE_LIGHT
    return p

def add_h3(doc, text):
    p = doc.add_paragraph()
    add_paragraph_spacing(p, before=120, after=60)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11.5)
    run.font.bold = True
    run.font.color.rgb = BLUE
    return p

def add_para(doc, text, bold=False, size=10.5, color=DARK, before=40, after=40, indent_left=None):
    p = doc.add_paragraph()
    add_paragraph_spacing(p, before=before, after=after)
    if indent_left:
        pPr = p._p.get_or_add_pPr()
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str(indent_left))
        pPr.append(ind)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    return p

def add_bullet(doc, text, level=0, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    add_paragraph_spacing(p, before=20, after=20)
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.5)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = DARK
    return p

def add_step(doc, text, size=10.5):
    p = doc.add_paragraph(style='List Number')
    add_paragraph_spacing(p, before=30, after=30)
    p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = DARK
    return p

def add_space(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        add_paragraph_spacing(p, before=0, after=0)

def add_tip(doc, text, bg=BLUE_BG, icon="💡", border_color="2E75B6", text_color=None):
    tbl = doc.add_table(rows=1, cols=2)
    set_table_width(tbl, 9026)
    row = tbl.rows[0]
    icon_c = row.cells[0]
    text_c = row.cells[1]
    styled_cell(icon_c, icon, bg=bg, font_size=11,
                align=WD_ALIGN_PARAGRAPH.CENTER, col_width=600)
    set_cell_borders(icon_c, border_color)
    styled_cell(text_c, text, bg=bg, font_size=9.5,
                font_color=RGBColor(0x1E, 0x3A, 0x5F) if text_color is None else text_color,
                col_width=8426)
    set_cell_borders(text_c, border_color)
    add_space(doc)
    return tbl

def add_warning(doc, text):
    return add_tip(doc, text, bg=ORANGE_BG, icon="⚠", border_color="B45309",
                   text_color=ORANGE)

def add_info_table(doc, rows_data, col_widths=(3000, 6026)):
    tbl = doc.add_table(rows=len(rows_data), cols=2)
    set_table_width(tbl, sum(col_widths))
    for i, (label, value) in enumerate(rows_data):
        row = tbl.rows[i]
        grey_label_cell(row.cells[0], label, col_width=col_widths[0])
        value_cell(row.cells[1], value, col_width=col_widths[1])
    return tbl

def add_header_table(doc, headers, rows_data, col_widths):
    tbl = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    set_table_width(tbl, sum(col_widths))
    # Header row
    for i, (h, w) in enumerate(zip(headers, col_widths)):
        header_cell(tbl.rows[0].cells[i], h, col_width=w)
    # Data rows
    for ri, row_data in enumerate(rows_data):
        for ci, (val, w, opts) in enumerate(zip(row_data, col_widths, [{}]*len(row_data))):
            if isinstance(val, tuple):
                txt, is_label = val
                if is_label:
                    grey_label_cell(tbl.rows[ri+1].cells[ci], txt, col_width=w)
                else:
                    value_cell(tbl.rows[ri+1].cells[ci], txt, col_width=w)
            else:
                value_cell(tbl.rows[ri+1].cells[ci], val, col_width=w)
    return tbl

def add_page_break(doc):
    doc.add_page_break()

# ──────────────────────────────────────────────────────────────────────────────
# PAGE DE COUVERTURE
# ──────────────────────────────────────────────────────────────────────────────

# Titre principal
add_space(doc, 4)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_paragraph_spacing(p, before=0, after=20)
run = p.add_run("SECURIONIS CHANTIERS")
run.font.name = "Arial"
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = BLUE

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_paragraph_spacing(p2, before=0, after=60)
run2 = p2.add_run("Application d'inspection SST")
run2.font.name = "Arial"
run2.font.size = Pt(14)
run2.font.color.rgb = BLUE_LIGHT

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_paragraph_spacing(p3, before=60, after=20)
run3 = p3.add_run("MODE D'EMPLOI")
run3.font.name = "Arial"
run3.font.size = Pt(22)
run3.font.bold = True
run3.font.color.rgb = BLUE

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_paragraph_spacing(p4, before=0, after=80)
run4 = p4.add_run("Guide complet d'utilisation — Version 1.1, Avril 2026")
run4.font.name = "Arial"
run4.font.size = Pt(12)
run4.font.color.rgb = DARK

# Tableau info couverture
add_space(doc, 2)
cover_tbl = doc.add_table(rows=3, cols=2)
set_table_width(cover_tbl, 7000)
cover_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
cw = [2500, 4500]
grey_label_cell(cover_tbl.rows[0].cells[0], "Version", col_width=cw[0])
value_cell(cover_tbl.rows[0].cells[1], "1.1 — Avril 2026", col_width=cw[1])
grey_label_cell(cover_tbl.rows[1].cells[0], "URL", col_width=cw[0])
value_cell(cover_tbl.rows[1].cells[1], "https://chantiers.securionis.com", col_width=cw[1])
grey_label_cell(cover_tbl.rows[2].cells[0], "Destinataires", col_width=cw[0])
value_cell(cover_tbl.rows[2].cells[1], "Inspecteurs SST et Administrateurs", col_width=cw[1])

add_page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# 1. PRÉSENTATION
# ──────────────────────────────────────────────────────────────────────────────
add_h1(doc, "1. Présentation")
add_para(doc, "Securionis Chantiers est une application web dédiée aux inspecteurs SST pour réaliser des contrôles de conformité sur les chantiers de construction en Suisse. Elle couvre l'ensemble du cycle d'inspection : préparation, visite terrain, documentation, rapport PDF et suivi des non-conformités.")
add_space(doc)
add_h2(doc, "Fonctionnalités principales")
for txt in [
    "Gestion des chantiers (création, modification, archivage)",
    "Visites d'inspection avec checklist réglementaire (447 points SUVA)",
    "Évaluation par point : Conforme, Non-conforme, Remarques, Pas nécessaire",
    "Prise de photos annotées sur le terrain (max. 10 par point)",
    "Analyse IA des photos et assistant juridique (références suisses)",
    "Génération de rapports PDF avec logo et signature",
    "Envoi des rapports par email à la liste des destinataires",
    "Suivi des écarts (non-conformités) et de leur correction",
    "Export Excel global ou par chantier",
    "Mode hors-ligne (PWA) avec synchronisation automatique",
    "Administration : utilisateurs, entreprise, points de contrôle",
]:
    add_bullet(doc, txt)
add_space(doc)
add_h2(doc, "Accès")
add_info_table(doc, [
    ("URL", "https://chantiers.securionis.com"),
    ("Navigateurs", "Chrome, Firefox, Safari, Edge (versions récentes)"),
    ("Appareils", "PC, Mac, tablette, smartphone (Android & iOS)"),
], col_widths=(2500, 6526))
add_space(doc)
add_tip(doc, "Pour une utilisation optimale sur le terrain, installez l'application sur votre tablette (voir section 11 — Mode hors-ligne).")

add_page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# 2. CONNEXION
# ──────────────────────────────────────────────────────────────────────────────
add_h1(doc, "2. Connexion et compte")
add_h3(doc, "2.1 Se connecter")
for txt in [
    "Ouvrez https://chantiers.securionis.com dans votre navigateur",
    "Saisissez votre adresse email et votre mot de passe",
    "Cliquez sur Se connecter",
]:
    add_step(doc, txt)
add_space(doc)
add_tip(doc, "La session dure 7 jours. Vous restez connecté même si vous fermez le navigateur.")
add_space(doc)
add_h3(doc, "2.2 Mot de passe oublié")
for txt in [
    "Sur la page de connexion, cliquez sur Mot de passe oublié ?",
    "Saisissez votre adresse email et cliquez sur Envoyer le lien",
    "Ouvrez l'email reçu et cliquez sur le lien de réinitialisation",
    "Choisissez un nouveau mot de passe (min. 8 caractères, majuscule, minuscule, chiffre)",
    "Cliquez sur Réinitialiser le mot de passe",
]:
    add_step(doc, txt)
add_space(doc)
add_h3(doc, "2.3 Rôles utilisateurs")
tbl = doc.add_table(rows=3, cols=2)
set_table_width(tbl, 9026)
header_cell(tbl.rows[0].cells[0], "Rôle", col_width=2000)
header_cell(tbl.rows[0].cells[1], "Droits", col_width=7026)
grey_label_cell(tbl.rows[1].cells[0], "Inspecteur", col_width=2000)
value_cell(tbl.rows[1].cells[1], "Créer et gérer ses propres chantiers, réaliser des visites, consulter ses rapports", col_width=7026)
grey_label_cell(tbl.rows[2].cells[0], "Administrateur", col_width=2000)
value_cell(tbl.rows[2].cells[1], "Accès complet : gestion des utilisateurs, entreprise, points de contrôle, tous les chantiers", col_width=7026)

add_page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# 3. DASHBOARD
# ──────────────────────────────────────────────────────────────────────────────
add_h1(doc, "3. Tableau de bord (Dashboard)")
add_para(doc, "Page d'accueil après connexion. Vue synthétique de l'activité de l'inspecteur connecté.")
add_space(doc)
add_h3(doc, "3.1 Indicateurs clés")
tbl = doc.add_table(rows=5, cols=2)
set_table_width(tbl, 9026)
header_cell(tbl.rows[0].cells[0], "Indicateur", col_width=3000)
header_cell(tbl.rows[0].cells[1], "Description", col_width=6026)
for label, val in [
    ("Chantiers actifs", "Nombre de chantiers non archivés"),
    ("NC ouvertes", "Non-conformités non encore corrigées"),
    ("Visites ce mois", "Nombre de visites terminées dans le mois en cours"),
    ("Taux de conformité", "Moyenne des points conformes sur les 3 derniers mois"),
]:
    ri = ["Chantiers actifs","NC ouvertes","Visites ce mois","Taux de conformité"].index(label) + 1
    grey_label_cell(tbl.rows[ri].cells[0], label, col_width=3000)
    value_cell(tbl.rows[ri].cells[1], val, col_width=6026)
add_space(doc)
add_h3(doc, "3.2 Graphique des non-conformités")
add_para(doc, "Évolution des NC sur les 6 derniers mois : barres bleues = NC ouvertes, barres vertes = NC corrigées.")
add_h3(doc, "3.3 Chantiers urgents")
add_para(doc, "Liste les chantiers avec NC dont le délai de correction est dépassé. Cliquez pour accéder directement au chantier.")
add_h3(doc, "3.4 Visites du mois")
add_para(doc, "Liste des visites réalisées dans le mois. Chaque ligne est cliquable pour accéder au rapport.")
add_h3(doc, "3.5 Export Excel global")
add_para(doc, "Le bouton Export Excel génère un fichier avec 4 feuilles : Chantiers, Visites, Écarts NC, Statistiques.")

add_page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# 4. CHANTIERS
# ──────────────────────────────────────────────────────────────────────────────
add_h1(doc, "4. Gestion des chantiers")
add_h3(doc, "4.1 Liste des chantiers")
for txt in [
    "Depuis la navigation, cliquez sur Chantiers pour afficher la liste des chantiers actifs",
    "Chaque carte : nom, adresse, nature des travaux, NC ouvertes, date dernière visite",
    "Barre de recherche : filtre par nom, adresse ou nature des travaux",
    "Onglet Archives : chantiers archivés accessibles en lecture",
]:
    add_bullet(doc, txt)
add_space(doc)
add_h3(doc, "4.2 Créer un nouveau chantier")
for txt in [
    "Cliquez sur Nouveau chantier",
    "Adresse et Nature des travaux sont obligatoires",
    "Optionnel : Nom du chantier, N° CAMAC, N° parcelle, N° ECA, Réf. communale, Contact sur site",
    "Cliquez sur Créer le chantier",
]:
    add_step(doc, txt)
add_space(doc)
add_h3(doc, "4.3 Page détail d'un chantier")
tbl = doc.add_table(rows=7, cols=2)
set_table_width(tbl, 9026)
header_cell(tbl.rows[0].cells[0], "Section", col_width=2800)
header_cell(tbl.rows[0].cells[1], "Contenu", col_width=6226)
for ri, (sec, cont) in enumerate([
    ("Informations", "Données du chantier, bouton Modifier, bouton Archive"),
    ("Documents", "Permis, plans, rapports classés par catégorie"),
    ("Destinataires", "Personnes recevant les rapports par email"),
    ("Visites", "Timeline avec statut, NC et bouton suppression pour visites non terminées"),
    ("Comparaison N/N-1", "Évolution des NC entre les deux dernières visites"),
    ("Non-conformités", "Liste complète des écarts avec statut et délai"),
], 1):
    grey_label_cell(tbl.rows[ri].cells[0], sec, col_width=2800)
    value_cell(tbl.rows[ri].cells[1], cont, col_width=6226)
add_space(doc)
add_h3(doc, "4.4 Archiver / Désarchiver")
add_para(doc, "Bouton Archive (icône boîte) sur la page chantier. Le chantier n'apparaît plus dans la liste principale mais reste accessible via l'onglet Archives.")
add_warning(doc, "Un chantier archivé ne peut plus recevoir de nouvelles visites.")
add_space(doc)
add_h3(doc, "4.5 Gérer les destinataires")
add_para(doc, "Section Destinataires sur la page chantier. Cliquez + Ajouter, saisissez Nom, Organisation (optionnel), Email, puis Enregistrer.")
add_tip(doc, "Vous pouvez ajouter plusieurs destinataires. Tous recevront le rapport en un seul envoi.")

add_page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# 5. DOCUMENTS
# ──────────────────────────────────────────────────────────────────────────────
add_h1(doc, "5. Documents par chantier")
add_para(doc, "Chaque chantier dispose d'un espace documentaire pour centraliser les pièces administratives et techniques.")
add_space(doc)
add_h3(doc, "5.1 Catégories disponibles")
for cat in ["Permis de construire", "Plans", "Rapport ECA", "Autorisation travaux dangereux", "Certificat entreprise", "Autre"]:
    add_bullet(doc, cat)
add_space(doc)
add_h3(doc, "5.2 Ajouter un document")
for txt in [
    "Cliquez sur Ajouter un document dans la section Documents",
    "Saisissez le nom du document",
    "Choisissez la catégorie",
    "Ajoutez une description (optionnel)",
    "Sélectionnez le fichier : PDF, Word, Excel, Image (max. 50 Mo)",
    "Cliquez sur Enregistrer",
]:
    add_step(doc, txt)
add_space(doc)
add_h3(doc, "5.3 Nouvelle version")
add_para(doc, "Cliquez sur Nouvelle version pour remplacer un fichier sans perdre l'historique. La version s'incrémente automatiquement (v2, v3...).")

add_page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# 6. VISITES
# ──────────────────────────────────────────────────────────────────────────────
add_h1(doc, "6. Réaliser une visite d'inspection")

add_h3(doc, "6.1 Préparer une visite")
add_para(doc, "La préparation permet de configurer les catégories et thèmes à inspecter avant de se rendre sur le chantier.")
for txt in [
    "Cliquez sur Préparer la visite depuis la page chantier",
    "Cochez les catégories à contrôler (ex. Accès & Sols, Échafaudages, Électricité...)",
    "Utilisez la barre de recherche pour trouver rapidement une catégorie ou un thème",
    "Affinez en sélectionnant les thèmes (Tout cocher / Tout décocher disponible)",
    "La préparation est sauvegardée pour utilisation lors de la visite",
]:
    add_step(doc, txt)

add_space(doc)
add_h3(doc, "6.2 Démarrer une nouvelle visite")
for txt in [
    "Sur la page chantier, cliquez sur Nouvelle visite",
    "Sélectionnez les catégories et thèmes à inspecter",
    "Cliquez sur Démarrer la visite",
]:
    add_step(doc, txt)

add_space(doc)
add_h3(doc, "6.3 Champs d'en-tête de la visite")
tbl = doc.add_table(rows=3, cols=2)
set_table_width(tbl, 9026)
header_cell(tbl.rows[0].cells[0], "Champ", col_width=3000)
header_cell(tbl.rows[0].cells[1], "Description", col_width=6026)
grey_label_cell(tbl.rows[1].cells[0], "Renseignements donnés par", col_width=3000)
value_cell(tbl.rows[1].cells[1], "Nom de la personne présente sur site lors de la visite", col_width=6026)
grey_label_cell(tbl.rows[2].cells[0], "Remarques générales", col_width=3000)
value_cell(tbl.rows[2].cells[1], "Observations globales sur la visite, apparaissent dans une section dédiée du rapport PDF", col_width=6026)

add_space(doc)
add_h3(doc, "6.4 Évaluer les points de contrôle")
add_para(doc, "Pour chaque point de contrôle, cliquez sur l'un des 4 boutons :")
add_space(doc)
tbl = doc.add_table(rows=5, cols=2)
set_table_width(tbl, 9026)
header_cell(tbl.rows[0].cells[0], "Bouton", col_width=2000)
header_cell(tbl.rows[0].cells[1], "Signification et effet dans le rapport", col_width=7026)
badge_cell(tbl.rows[1].cells[0], "Conforme", GREEN_BG, GREEN, col_width=2000)
value_cell(tbl.rows[1].cells[1], "Le point respecte la réglementation. Apparaît dans les statistiques.", col_width=7026)
badge_cell(tbl.rows[2].cells[0], "Non-conforme", RED_BG, RED, col_width=2000)
value_cell(tbl.rows[2].cells[1], "Un écart est constaté. Crée un écart avec délai de correction. Apparaît dans la section Constatations du PDF.", col_width=7026)
badge_cell(tbl.rows[3].cells[0], "Remarques", ORANGE_BG, ORANGE, col_width=2000)
value_cell(tbl.rows[3].cells[1], "Observation à signaler sans caractère d'infraction. Apparaît dans la section Remarques du PDF (fond ambre).", col_width=7026)
badge_cell(tbl.rows[4].cells[0], "Pas nécessaire", GREY_BG, RGBColor(0x55,0x55,0x55), col_width=2000)
value_cell(tbl.rows[4].cells[1], "Le point ne s'applique pas à ce chantier. Exclu des statistiques.", col_width=7026)

add_space(doc)
add_h3(doc, "6.5 Ajouter une remarque textuelle")
add_para(doc, "Sous les boutons de réponse, un champ texte Remarque permet de détailler la constatation. Ce texte apparaît dans le rapport PDF sous le point concerné.")

add_space(doc)
add_h3(doc, "6.6 Prendre des photos")
for txt in [
    "Cliquez sur Appareil photo ou Galerie (max. 10 photos par point de contrôle)",
    "L'éditeur d'annotation s'ouvre automatiquement",
    "Annotez : flèche, cercle, texte, dessin libre — 5 couleurs, 3 épaisseurs",
    "Cliquez sur Valider pour sauvegarder l'annotation",
    "Pour modifier : survolez la photo et cliquez sur l'icône crayon",
]:
    add_step(doc, txt)

add_space(doc)
add_h3(doc, "6.7 Analyse IA des photos")
add_para(doc, "Disponible dès qu'une photo est uploadée sur un point de contrôle.")
for txt in [
    "Cliquez sur Analyse IA",
    "L'IA analyse la dernière photo et détecte les dangers potentiels",
    "Suggestions avec niveau de sévérité : critique / majeur / mineur",
    "Cliquez sur Appliquer la remarque ou Appliquer la conformité",
]:
    add_step(doc, txt)
add_space(doc)
add_tip(doc, "L'analyse IA est particulièrement utile pour détecter automatiquement des risques visibles sur photo : équipements manquants, zones non sécurisées, etc.")

add_space(doc)
add_h3(doc, "6.8 Assistant juridique")
add_para(doc, "Répond aux questions réglementaires pour chaque point de contrôle.")
for txt in [
    "Cliquez sur Assistant juridique pour déployer le panneau",
    "Choisissez une question rapide ou tapez votre propre question",
    "L'IA répond avec les références légales suisses (OTConst, SUVA, SIA...)",
    "Cliquez sur Copier dans la remarque pour insérer la réponse",
]:
    add_step(doc, txt)

add_space(doc)
add_h3(doc, "6.9 Ajouter des catégories en cours de visite")
add_para(doc, "Il est possible d'ajouter de nouvelles catégories/thèmes sans perdre les réponses déjà saisies : faites défiler jusqu'en bas et cliquez sur + Catégories / Thèmes.")

add_space(doc)
add_h3(doc, "6.10 Sauvegarde automatique")
add_para(doc, "Chaque modification est sauvegardée automatiquement toutes les 2 secondes. En mode hors-ligne, les données sont stockées localement et synchronisées au retour du réseau.")

add_space(doc)
add_h3(doc, "6.11 Valider la visite")
for txt in [
    "Cliquez sur Valider la visite en bas de la checklist",
    "Pour chaque point Non-conforme, saisissez un délai de correction (ex. « Immédiatement », « 7 jours », « 30.06.2026 »)",
    "Cliquez sur Suivant ou Valider et terminer",
    "La visite passe au statut Terminée — vous êtes redirigé vers le rapport",
]:
    add_step(doc, txt)

add_space(doc)
add_h3(doc, "6.12 Supprimer une visite en cours")
add_para(doc, "Uniquement pour les visites au statut En cours ou Brouillon (pas les visites Terminées).")
for txt in [
    "Sur la page chantier, dans la timeline, cliquez sur l'icône corbeille",
    "Confirmez avec Oui dans la fenêtre de confirmation",
]:
    add_step(doc, txt)
add_warning(doc, "La suppression est définitive et irréversible. Toutes les réponses et photos associées sont effacées.")

add_page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# 7. RAPPORT
# ──────────────────────────────────────────────────────────────────────────────
add_h1(doc, "7. Rapport de visite")

add_h3(doc, "7.1 Contenu du rapport PDF")
tbl = doc.add_table(rows=9, cols=2)
set_table_width(tbl, 9026)
header_cell(tbl.rows[0].cells[0], "Section PDF", col_width=3000)
header_cell(tbl.rows[0].cells[1], "Contenu", col_width=6026)
for ri, (sec, cont) in enumerate([
    ("En-tête", "Logo de l'entreprise, titre Rapport de visite"),
    ("Informations", "Inspecteur, date, chantier, adresse, nature des travaux, références"),
    ("Renseignements par", "Nom de la personne rencontrée sur site"),
    ("Remarques générales", "Observations globales de la visite (si renseignées)"),
    ("Constatations", "Points Non-conformes : intitulé, remarque, photos, délai, statut"),
    ("Remarques", "Points Remarques : intitulé, remarque, photos (fond ambre)"),
    ("Signature", "Signature de l'inspecteur (si configurée)"),
    ("Copie(s)", "Liste des destinataires + pied de page entreprise"),
], 1):
    grey_label_cell(tbl.rows[ri].cells[0], sec, col_width=3000)
    value_cell(tbl.rows[ri].cells[1], cont, col_width=6026)

add_space(doc)
add_h3(doc, "7.2 Générer le PDF")
for txt in [
    "Sur la page rapport, cliquez sur Télécharger PDF",
    "Le PDF est généré et sauvegardé dans le cloud",
    "Cliquez sur Voir le rapport pour l'ouvrir dans l'aperçu",
]:
    add_step(doc, txt)

add_space(doc)
add_h3(doc, "7.3 Envoyer par email")
for txt in [
    "Assurez-vous qu'au moins un destinataire est configuré pour le chantier",
    "Cliquez sur Envoyer par email",
    "Le rapport est envoyé à tous les destinataires en un seul email avec signature de l'entreprise",
]:
    add_step(doc, txt)
add_tip(doc, "Après correction de toutes les NC, un bandeau vert vous invite à régénérer le rapport et le renvoyer.")

add_page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# 8. ÉCARTS
# ──────────────────────────────────────────────────────────────────────────────
add_h1(doc, "8. Suivi des non-conformités (Écarts)")

add_h3(doc, "8.1 Statuts des écarts")
tbl = doc.add_table(rows=4, cols=2)
set_table_width(tbl, 9026)
header_cell(tbl.rows[0].cells[0], "Statut", col_width=2500)
header_cell(tbl.rows[0].cells[1], "Signification", col_width=6526)
badge_cell(tbl.rows[1].cells[0], "Ouvert", RED_BG, RED, col_width=2500)
value_cell(tbl.rows[1].cells[1], "Écart constaté, aucune action corrective engagée", col_width=6526)
badge_cell(tbl.rows[2].cells[0], "En cours de correction", ORANGE_BG, ORANGE, col_width=2500)
value_cell(tbl.rows[2].cells[1], "Des mesures correctives ont été mises en place", col_width=6526)
badge_cell(tbl.rows[3].cells[0], "Corrigé", GREEN_BG, GREEN, col_width=2500)
value_cell(tbl.rows[3].cells[1], "L'écart a été entièrement résolu", col_width=6526)

add_space(doc)
add_h3(doc, "8.2 Mettre à jour un statut")
for txt in [
    "Sur la page chantier, section Non-conformités, cliquez sur le badge de statut",
    "Sélectionnez le nouveau statut (Ouvert → En cours / Corrigé, En cours → Corrigé)",
    "La mise à jour est immédiate",
]:
    add_step(doc, txt)

add_space(doc)
add_h3(doc, "8.3 Comparaison entre visites (N/N-1)")
add_para(doc, "Compare automatiquement les deux dernières visites terminées :")
add_bullet(doc, "Corrigée (vert) : point Non-conforme en N-1, Conforme en N")
add_bullet(doc, "Persistante (orange) : point Non-conforme lors des deux visites")
add_bullet(doc, "Nouvelle (rouge) : point Non-conforme en N, Conforme en N-1")
add_tip(doc, "Cette comparaison est disponible à partir de la 2e visite sur un chantier.")

add_page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# 9. EXPORT EXCEL
# ──────────────────────────────────────────────────────────────────────────────
add_h1(doc, "9. Export Excel")
add_h3(doc, "9.1 Export global (Dashboard)")
add_para(doc, "4 feuilles : Chantiers, Visites, Écarts NC, Statistiques.")

add_h3(doc, "9.2 Export par chantier (page chantier)")
add_para(doc, "4 feuilles : Chantier, Visites, Écarts NC, Réponses détaillées (chaque point avec valeur, remarque et base légale).")
add_para(doc, "Cliquez sur l'icône Excel (bouton vert) sur la page concernée. Le fichier se télécharge automatiquement.")

add_page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# 10. ADMINISTRATION
# ──────────────────────────────────────────────────────────────────────────────
add_h1(doc, "10. Administration")
add_para(doc, "Accessible uniquement aux Administrateurs via la navigation (icône engrenage).")
add_space(doc)

add_h3(doc, "10.1 Points de contrôle (Admin > Points de contrôle)")
add_para(doc, "Base de 447 points SUVA (26 catégories, 442 thèmes) :")
for txt in [
    "Filtrer par catégorie, thème, statut ou recherche texte libre",
    "Activer / Désactiver un point (les inactifs n'apparaissent plus dans les nouvelles visites)",
    "Modifier un point existant : cliquez dessus pour ouvrir le formulaire",
    "Créer un point : bouton + Nouveau point — remplissez intitulé, catégorie, thème, critère, base légale, explications, documents PDF associés (max. 5)",
]:
    add_bullet(doc, txt)

add_space(doc)
add_h3(doc, "10.2 Utilisateurs (Admin > Utilisateurs)")
for txt in [
    "Voir la liste avec rôles",
    "Créer un utilisateur : + Nouvel utilisateur, saisissez nom, email, rôle",
    "Modifier le rôle d'un utilisateur existant",
    "Supprimer un utilisateur",
]:
    add_bullet(doc, txt)
add_warning(doc, "Le rôle d'un utilisateur ne peut pas être modifié par lui-même (protection de sécurité).")

add_space(doc)
add_h3(doc, "10.3 Entreprise (Admin > Entreprise)")
add_para(doc, "Configurez les informations qui apparaissent dans les rapports et emails : Nom, Adresse, NPA, Ville, Téléphone, Email, Logo (PNG ou JPEG recommandé).")

add_page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# 11. PWA / HORS-LIGNE
# ──────────────────────────────────────────────────────────────────────────────
add_h1(doc, "11. Mode hors-ligne (PWA)")
add_para(doc, "L'application fonctionne sans connexion internet, idéale pour les visites en sous-sol ou en zone sans réseau.")
add_space(doc)

add_h3(doc, "11.1 Installer l'application")
for txt in [
    "Ouvrez https://chantiers.securionis.com dans Chrome (Android) ou Safari (iOS)",
    "Appuyez sur le menu navigateur → Ajouter à l'écran d'accueil ou Installer l'application",
    "L'application s'ouvre en plein écran comme une app native",
]:
    add_step(doc, txt)

add_space(doc)
add_h3(doc, "11.2 Fonctionnement hors-ligne")
add_bullet(doc, "Bandeau rouge Hors-ligne quand la connexion est perdue")
add_bullet(doc, "Saisies sauvegardées localement (IndexedDB)")
add_bullet(doc, "Photos stockées en attente de synchronisation")
add_bullet(doc, "Checklist entièrement utilisable sans réseau")

add_space(doc)
add_h3(doc, "11.3 Synchronisation au retour du réseau")
add_bullet(doc, "Bandeau orange « X modifications en attente » s'affiche")
add_bullet(doc, "Synchronisation automatique au retour du réseau")
add_bullet(doc, "Ou cliquez sur Synchroniser pour forcer la synchronisation immédiate")
add_tip(doc, "Pour les visites longues, activez le mode avion après avoir chargé la checklist. Toutes les saisies seront sauvegardées localement.")

add_page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# 12. RACCOURCIS
# ──────────────────────────────────────────────────────────────────────────────
add_h1(doc, "12. Raccourcis et astuces")
shortcuts = [
    ("Rechercher un chantier", "Barre de recherche, page Chantiers"),
    ("Rechercher une catégorie/thème", "Barre de recherche lors de la création de visite ou en cours"),
    ("Annoter une photo existante", "Survolez la photo → icône crayon"),
    ("Copier une suggestion IA", "Bouton Copier dans la remarque sous la suggestion"),
    ("Ajouter un thème en cours", "+ Catégories / Thèmes en bas de checklist"),
    ("Supprimer une visite", "Icône corbeille dans la timeline (visites non terminées uniquement)"),
    ("Voir le rapport", "Cliquez sur la visite Terminée dans la timeline"),
    ("Régénérer un rapport", "Page rapport → Télécharger PDF (écrase l'ancien)"),
    ("Archiver un chantier", "Bouton Archive (boîte) sur la page chantier"),
    ("Export Excel chantier", "Icône Excel (vert) sur la page chantier"),
    ("Export Excel global", "Bouton Export Excel sur le Dashboard"),
]
tbl = doc.add_table(rows=1 + len(shortcuts), cols=2)
set_table_width(tbl, 9026)
header_cell(tbl.rows[0].cells[0], "Action", col_width=4000)
header_cell(tbl.rows[0].cells[1], "Comment faire", col_width=5026)
for ri, (action, how) in enumerate(shortcuts, 1):
    grey_label_cell(tbl.rows[ri].cells[0], action, col_width=4000)
    value_cell(tbl.rows[ri].cells[1], how, col_width=5026)

add_page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# 13. SÉCURITÉ
# ──────────────────────────────────────────────────────────────────────────────
add_h1(doc, "13. Sécurité et confidentialité")
for txt in [
    "Authentification sécurisée avec session de 7 jours",
    "Toutes les communications chiffrées (HTTPS/TLS)",
    "Données isolées par entreprise (accès impossible aux données d'une autre entreprise)",
    "Photos et rapports dans un espace sécurisé avec accès contrôlé",
    "Journal d'audit des actions sensibles (envois email, changements de statut, etc.)",
    "Chaque utilisateur n'accède qu'aux chantiers qui lui sont assignés",
]:
    add_bullet(doc, txt)
add_space(doc)
add_warning(doc, "Ne partagez jamais vos identifiants de connexion. En cas de suspicion de compromission, changez immédiatement votre mot de passe via la fonction Mot de passe oublié.")

add_page_break(doc)

# ──────────────────────────────────────────────────────────────────────────────
# 14. SUPPORT
# ──────────────────────────────────────────────────────────────────────────────
add_h1(doc, "14. Support")
add_info_table(doc, [
    ("Support technique", "Contactez votre administrateur ou l'adresse email configurée dans Admin > Entreprise"),
    ("URL application", "https://chantiers.securionis.com"),
    ("Version", "1.1 — Avril 2026"),
], col_widths=(2800, 6226))
add_space(doc, 3)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_paragraph_spacing(p, before=200, after=60)
run = p.add_run("Securionis Chantiers — Santé et Sécurité au Travail")
run.font.name = "Arial"
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("https://chantiers.securionis.com")
run2.font.name = "Arial"
run2.font.size = Pt(9)
run2.font.color.rgb = BLUE_LIGHT

# ──────────────────────────────────────────────────────────────────────────────
# SAUVEGARDE
# ──────────────────────────────────────────────────────────────────────────────
out_path = "docs/Mode_Emploi_Securionis_Chantiers.docx"
doc.save(out_path)
print(f"Document genere : {out_path}")
